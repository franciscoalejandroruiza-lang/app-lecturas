"""
fill_formats.py
Lógica principal para:
 1) Leer el catálogo maestro y el Excel de consumo mensual (hoja OK).
 2) Ubicar, para cada número de serie, el bloque de 5 columnas del mes elegido.
 3) Llenar los formatos Word originales (uno por ubicación) con las lecturas
    del mes (inicial y/o final) y exportarlos a PDF.

Los formatos Word originales se usan como plantilla viva: se localizan las
tablas de cada equipo por su número de serie y solo se editan las celdas de
"Lectura inicial", "Lectura final" y "Volumen mensual", preservando el resto
del documento (logos, formato, firmas, etc.) exactamente igual.
"""
import re
import copy
import subprocess
import tempfile
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

# ---------- Orden de los 5 (o 6) renglones de lectura en cada formato ----------
LECTURA_ROW_LABELS = [
    ("carta_bn", "COPIA O IMPRESIÓN EN BLANCO Y NEGRO TAMAÑO CARTA"),
    ("oficio_bn", "COPIA O IMPRESIÓN EN BLANCO Y NEGRO TAMAÑO OFICIO"),
    ("doble_carta_bn", "COPIA O IMPRESIÓN EN BLANCO Y NEGRO TAMAÑO DOBLE CARTA"),
    ("carta_color", "COPIA O IMPRESIÓN EN COLOR TAMAÑO CARTA"),
    ("oficio_color", "COPIA O IMPRESIÓN EN COLOR TAMAÑO OFICIO"),
    ("digitalizacion", "POR HOJA DE DIGITALIZACIÓN"),
]

# Orden de los 5 campos tal como aparecen en cada bloque mensual del Excel "OK"
EXCEL_MES_CAMPOS = ["carta_bn", "oficio_bn", "carta_color", "oficio_color", "digitalizacion"]


def classify(t):
    txt = ' | '.join(c.text for r in t.rows for c in r.cells)
    if 'FORMATO PARA' in txt or ('Estado' in txt and 'CLUE' in txt):
        return 'A_header'
    if 'Unidad Médica o Administrativa' in txt:
        return 'B_unidad'
    if 'NÚMERO DE' in txt.upper() and 'SERIE' in txt.upper():
        return 'C_equipo'
    if 'Lectura inicial' in txt:
        return 'D_lectura'
    if 'FIRMAS' in txt.upper() or 'Nombre y Firma' in txt:
        return 'E_firmas'
    return 'X_unknown'


def get_serie_from_C(t):
    for r in t.rows:
        vals = [c.text.strip() for c in r.cells]
        if len(vals) < 3:
            continue
        if vals[1].upper() in ('NÚMERO DE SERIE', 'NUMERO DE SERIE'):
            data_vals = [v for v in vals[2:] if v]
            if data_vals:
                return data_vals[0]
    return None


def fmt_num(v):
    """Formatea un número con comas de miles, tal como aparece en los formatos originales."""
    if v is None or v == '':
        return ''
    try:
        f = float(v)
        if f == int(f):
            return f"{int(f):,}"
        return f"{f:,.2f}"
    except (ValueError, TypeError):
        return str(v)


def set_cell_text(cell, text):
    """Reemplaza el texto de una celda preservando la fuente del primer run existente."""
    para = cell.paragraphs[0]
    # Guardar formato del primer run si existe
    ref_run = para.runs[0] if para.runs else None
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    # limpiar párrafos extra
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    run = para.add_run(text)
    if ref_run is not None:
        run.font.name = ref_run.font.name
        run.font.size = ref_run.font.size
        run.font.bold = ref_run.font.bold
        rPr = ref_run._element.find(qn('w:rPr'))
        if rPr is not None:
            east_asian = rPr.find(qn('w:rFonts'))
            if east_asian is not None:
                run._element.get_or_add_rPr().append(copy.deepcopy(east_asian))


def fill_lectura_table(table, inicial, final_vals, modo):
    """
    inicial / final_vals: dicts con llaves carta_bn, oficio_bn, doble_carta_bn,
    carta_color, oficio_color, digitalizacion -> valor numérico o None.
    modo: 'inicial' (solo llena columna Lectura inicial) o 'final' (llena
    inicial + final + calcula volumen mensual).
    """
    row_idx = 1  # fila 0 es encabezado
    for key, _label in LECTURA_ROW_LABELS:
        if row_idx >= len(table.rows):
            break
        row = table.rows[row_idx]
        cells = row.cells
        val_ini = inicial.get(key) if inicial else None
        set_cell_text(cells[2], fmt_num(val_ini))
        if modo == 'final':
            val_fin = final_vals.get(key) if final_vals else None
            set_cell_text(cells[3], fmt_num(val_fin))
            vol = ''
            if val_ini not in (None, '') and val_fin not in (None, ''):
                try:
                    vol = fmt_num(float(val_fin) - float(val_ini))
                except (ValueError, TypeError):
                    vol = ''
            set_cell_text(cells[4], vol)
        else:
            set_cell_text(cells[3], '')
            set_cell_text(cells[4], '')
        row_idx += 1


def fill_fecha_hora(a_table, texto):
    """Escribe la fecha y hora de lectura en la celda de valor de la tabla de
    encabezado (A_header), preservando el resto de esa tabla intacto."""
    if not texto:
        return
    for row in a_table.rows:
        cells = row.cells
        texts = [c.text.strip() for c in cells]
        if any('Fecha y Hora de Lectura' in t for t in texts):
            seen = set()
            unique_cells = []
            for c in cells:
                if id(c._tc) not in seen:
                    seen.add(id(c._tc))
                    unique_cells.append(c)
            for c in unique_cells:
                t = c.text.strip()
                if t not in ('Fecha y Hora de Lectura', 'FORMATO PARA TOMA DE LECTURA'):
                    set_cell_text(c, texto)
                    return


def fill_template(template_path, readings_by_serie, modo, output_path):
    """
    readings_by_serie: { serie: {'inicial': {...}, 'final': {...}} }
    modo: 'inicial' o 'final'
    Guarda el docx resultante en output_path. Devuelve lista de series
    encontradas en la plantilla y lista de series del Excel que no se
    encontraron en la plantilla (para poder avisar al usuario).
    """
    doc = Document(str(template_path))
    tables = list(doc.tables)
    classes = [classify(t) for t in tables]
    n = len(tables)
    series_en_plantilla = []
    current_header = None
    i = 0
    while i < n:
        if classes[i] == 'A_header':
            current_header = tables[i]
        if classes[i] == 'C_equipo':
            serie = get_serie_from_C(tables[i])
            if serie:
                series_en_plantilla.append(serie)
                # la tabla de lectura es la siguiente D_lectura tras esta C_equipo
                j = i + 1
                if j < n and classes[j] == 'D_lectura':
                    r = readings_by_serie.get(serie.strip())
                    if r:
                        fill_lectura_table(tables[j], r.get('inicial', {}), r.get('final', {}), modo)
                        # La fecha y hora de lectura corresponde al momento en que se
                        # tomó la lectura FINAL (solo aplica en modo 'final'; en modo
                        # 'inicial' el formato va en blanco para que la operadora la
                        # anote a mano en sitio).
                        if modo == 'final' and current_header is not None:
                            fecha = (r.get('final') or {}).get('fecha_hora')
                            fill_fecha_hora(current_header, fecha)
        i += 1
    doc.save(str(output_path))
    faltantes = [s for s in readings_by_serie if s not in [x.strip() for x in series_en_plantilla]]
    return series_en_plantilla, faltantes


def _find_soffice():
    """Busca el ejecutable de LibreOffice. En Windows/Mac `soffice` casi nunca
    está en el PATH aunque LibreOffice sí esté instalado, así que se revisan
    también las rutas típicas de instalación."""
    import shutil as _shutil
    found = _shutil.which("soffice") or _shutil.which("soffice.exe")
    if found:
        return found
    candidatos = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/usr/bin/soffice",
        "/usr/local/bin/soffice",
    ]
    for c in candidatos:
        if Path(c).exists():
            return c
    raise RuntimeError(
        "No se encontró LibreOffice (soffice) instalado. Descárgalo de "
        "https://www.libreoffice.org/download/download/ e instálalo, o si "
        "ya lo instalaste, agrega la carpeta 'program' de LibreOffice al PATH."
    )


def docx_to_pdf(docx_path, out_dir):
    """Convierte un .docx a .pdf usando LibreOffice headless. Requiere LibreOffice instalado."""
    docx_path = str(docx_path)
    out_dir = str(out_dir)
    soffice = _find_soffice()
    result = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
        capture_output=True, text=True, timeout=120,
    )
    if result.returncode != 0:
        raise RuntimeError(f"Error convirtiendo a PDF: {result.stderr}")
    pdf_name = Path(docx_path).stem + ".pdf"
    return Path(out_dir) / pdf_name
